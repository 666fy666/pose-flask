"""
报告生成器模块
负责生成Word格式的分析报告
按照指定模板格式生成
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from typing import Dict, Any, Tuple

class ReportGenerator:
    """报告生成器类"""
    
    def __init__(self):
        """初始化报告生成器"""
        pass
    
    def generate_report(self, report_data: Dict[str, Any], output_dir: str, 
                       patient_name: str, shoulder_selection: str = 'left') -> str:
        """
        生成Word格式的分析报告
        
        Args:
            report_data: 报告数据
            output_dir: 输出目录
            patient_name: 患者姓名
            shoulder_selection: 肩部选择，'left'表示左肩，'right'表示右肩
            
        Returns:
            str: 报告文件路径
        """
        # 准备报告数据
        formatted_data = self._prepare_report_data(report_data, patient_name, output_dir, shoulder_selection)
        
        # 创建Word报告
        file_path, file_name = self._create_word_report(formatted_data, output_dir, patient_name, shoulder_selection)
        
        return file_path
    
    def _prepare_report_data(self, report_data: Dict[str, Any], patient_name: str, output_dir: str, shoulder_selection: str = 'left') -> Dict[str, Any]:
        """
        准备报告数据，按照模板格式组织
        
        Args:
            report_data: 原始报告数据
            patient_name: 患者姓名
            output_dir: 输出目录
            shoulder_selection: 肩部选择，'left'表示左肩，'right'表示右肩
            
        Returns:
            Dict[str, Any]: 格式化的报告数据
        """
        # 获取患者基本信息
        patient_info = report_data.get('patient_info', {})
        
        # 获取分析结果
        summary = report_data.get('summary', {})
        
        # 获取前视肩部数据（外展运动）
        front_shoulder_data = report_data.get('front_shoulder_data', {})
        front_left_data = front_shoulder_data.get('left_shoulder_data', {})
        front_right_data = front_shoulder_data.get('right_shoulder_data', {})
        
        # 获取侧视肩部数据（前屈运动）
        side_shoulder_data = report_data.get('side_shoulder_data', {})
        side_left_data = side_shoulder_data.get('left_shoulder_data', {})
        side_right_data = side_shoulder_data.get('right_shoulder_data', {})
        
        # 获取腕部数据
        wrist_data = report_data.get('wrist_data', {})
        
        # 准备基础信息
        base_info = {
            "name": patient_name,
            "age": patient_info.get('age', 0),
            "gender": patient_info.get('gender', ''),
            "height": patient_info.get('height', 0),
            "weight": patient_info.get('weight', 0)
        }
        
        # 准备前视左肩部角度数据
        front_left_shoulder_data = {
            "angle_speed": front_left_data.get('velocity_stages', [0, 0, 0, 0]),
            "max_angle": front_left_data.get('max_angle', 0),
            "max_angle_speed": front_left_data.get('max_velocity', 0)
        }
        
        # 准备前视右肩部角度数据
        front_right_shoulder_data = {
            "angle_speed": front_right_data.get('velocity_stages', [0, 0, 0, 0]),
            "max_angle": front_right_data.get('max_angle', 0),
            "max_angle_speed": front_right_data.get('max_velocity', 0)
        }
        
        # 根据肩部选择准备侧视数据
        if shoulder_selection == 'left':
            # 选择左肩，只准备左肩数据
            side_selected_shoulder_data = {
                "angle_speed": side_left_data.get('velocity_stages', [0, 0, 0, 0]),
                "max_angle": side_left_data.get('max_angle', 0),
                "max_angle_speed": side_left_data.get('max_velocity', 0),
                "shoulder_name": "左肩"
            }
        else:
            # 选择右肩，只准备右肩数据
            side_selected_shoulder_data = {
                "angle_speed": side_right_data.get('velocity_stages', [0, 0, 0, 0]),
                "max_angle": side_right_data.get('max_angle', 0),
                "max_angle_speed": side_right_data.get('max_velocity', 0),
                "shoulder_name": "右肩"
            }
        
        # 准备腕部关节数据
        wrist_joint_data = {
            "left_wrist_height_ratio": wrist_data.get('left_max_height', 0),
            "right_wrist_height_ratio": wrist_data.get('right_max_height', 0)
        }
        
        # 准备图片路径（使用相对路径）
        image_path = {
            "front_shoulder_angle_speed": os.path.join(output_dir, "..", "analysis_results", "正面_角度分析.png"),
            "side_shoulder_angle_speed": os.path.join(output_dir, "..", "analysis_results", "侧面_角度分析.png"),
            "back_wrist_height": os.path.join(output_dir, "..", "analysis_results", "背面_手腕高度.png"),
            # 关键帧图片路径
            "max_abduction_angles": os.path.join(output_dir, "..", "analysis_results", "max_abduction_angles.png"),
            "max_flexion_angle": os.path.join(output_dir, "..", "analysis_results", "max_flexion_angle.png"),
            "max_wrist_heights": os.path.join(output_dir, "..", "analysis_results", "max_wrist_heights.png")
        }
        
        # 检查关键帧图片是否存在，如果不存在则设置为空字符串
        for key in ["max_abduction_angles", "max_flexion_angle", "max_wrist_heights"]:
            if not os.path.exists(image_path[key]):
                image_path[key] = ""
        
        # 准备结果报告
        result_report = {
            "function_score": summary.get('function_score', 0),
            "function_report": summary.get('function_assessment', ''),
            "check_date": datetime.now().strftime("%Y-%m-%d")
        }
        
        # 返回格式化的数据
        return {
            "base_info": base_info,
            "front_left_shoulder_data": front_left_shoulder_data,
            "front_right_shoulder_data": front_right_shoulder_data,
            "side_selected_shoulder_data": side_selected_shoulder_data,
            "wrist_joint_data": wrist_joint_data,
            "image_path": image_path,
            "result_report": result_report
        }
    
    def _create_word_report(self, report_data: Dict[str, Any], output_dir: str, patient_name: str, shoulder_selection: str = 'left') -> Tuple[str, str]:
        """
        创建Word报告
        
        Args:
            report_data: 格式化的报告数据
            output_dir: 输出目录
            patient_name: 患者姓名
            shoulder_selection: 肩部选择，'left'表示左肩，'right'表示右肩
            
        Returns:
            Tuple[str, str]: (文件路径, 文件名)
        """
        # 创建一个新的 Word 文档
        doc = Document()

        # 添加文档标题，并设置居中
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        heading = doc.add_heading('肩关节运动报告', level=1)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 正确设置居中对齐

        # 添加段落
        doc.add_heading('一、一般信息', level=2)
        # 把信息放同一行表格中
        table = doc.add_table(rows=1, cols=10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '姓名:'
        hdr_cells[1].text = report_data['base_info']['name']
        hdr_cells[2].text = '年龄:'
        hdr_cells[3].text = str(report_data['base_info']['age'])
        hdr_cells[4].text = '性别:'
        hdr_cells[5].text = report_data['base_info']['gender']
        hdr_cells[6].text = '身高(cm):'
        hdr_cells[7].text = str(report_data['base_info']['height'])
        hdr_cells[8].text = '体重(kg):'
        hdr_cells[9].text = str(report_data['base_info']['weight'])

        # 添加段落
        doc.add_heading('二、前视外展运动信息', level=2)
        
        # 小标题
        doc.add_heading('1.平均角速度', level=3)
        # 添加表格，3行5列
        # 列名分别为：项目、0~45°、45~90°、90~135°、135~180°
        # 肩部角度数据list,用于填充表格
        
        temp_list = [['左肩平均角速度°/s'] + [str(x) for x in report_data['front_left_shoulder_data']['angle_speed']],
                     ['右肩平均角速度°/s'] + [str(x) for x in report_data['front_right_shoulder_data']['angle_speed']]]
    
        table = doc.add_table(rows=1, cols=5)  # 只创建标题行
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '角度范围'
        hdr_cells[1].text = '0~45°'
        hdr_cells[2].text = '45~90°'
        hdr_cells[3].text = '90~135°'
        hdr_cells[4].text = '135~180°'
        # 添加数据
        for i in range(len(temp_list)):
            row_cells = table.add_row().cells  # 为每个数据行添加一行
            row_cells[0].text = temp_list[i][0]
            row_cells[1].text = temp_list[i][1]
            row_cells[2].text = temp_list[i][2]
            row_cells[3].text = temp_list[i][3]
            row_cells[4].text = temp_list[i][4]


        # 小标题
        doc.add_heading('2.最大外展角度信息', level=3)
        # 创建3行3列的表格，第一行为表头
        table = doc.add_table(rows=3, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '角度信息'
        hdr_cells[1].text = '最大角度'
        hdr_cells[2].text = '最大角速度'
        # 左肩数据
        left_cells = table.rows[1].cells
        left_cells[0].text = '左肩'
        left_cells[1].text = str(report_data['front_left_shoulder_data']['max_angle'])
        left_cells[2].text = str(report_data['front_left_shoulder_data']['max_angle_speed'])
        # 右肩数据
        right_cells = table.rows[2].cells
        right_cells[0].text = '右肩'
        right_cells[1].text = str(report_data['front_right_shoulder_data']['max_angle'])
        right_cells[2].text = str(report_data['front_right_shoulder_data']['max_angle_speed'])

        # 小标题
        doc.add_heading('3.最大外展角角度视图', level=3)
        # 插入最大外展角角度视图
        if report_data['image_path']['max_abduction_angles'] and os.path.exists(report_data['image_path']['max_abduction_angles']):
            try:
                doc.add_picture(report_data['image_path']['max_abduction_angles'], width=Inches(6))
            except:
                doc.add_paragraph('最大外展角角度视图（图片加载失败）')
        else:
            doc.add_paragraph('最大外展角角度视图（图片未生成）')

        # 小标题
        doc.add_heading('4.角度和角度速度曲线图', level=3)
        # 插入图片,先不用操作
        # 添加图片
        try:
            doc.add_picture(report_data['image_path']['front_shoulder_angle_speed'], width=Inches(6))
        except:
            doc.add_paragraph('正面-肩部角度曲线图（图片加载失败）')

        # 直接插入分页符
        doc.add_page_break()

        
        # 小标题
        doc.add_heading('三、侧视前屈运动信息', level=2)
        # 角度范围	0~45°	45~90°	90~135°	135~180°
        # 平均角速度°/s	3.25	3.03	3.18	2.61
        doc.add_heading('1.平均角速度', level=3)
        
        # 获取选择的肩部数据
        selected_shoulder_data = report_data['side_selected_shoulder_data']
        shoulder_name = selected_shoulder_data['shoulder_name']
        
        # 准备侧视数据列表（只显示选择的肩部）
        side_temp_list = [[f'{shoulder_name}平均角速度°/s'] + [str(x) for x in selected_shoulder_data['angle_speed']]]
        
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '角度范围'
        hdr_cells[1].text = '0~45°'
        hdr_cells[2].text = '45~90°'
        hdr_cells[3].text = '90~135°'
        hdr_cells[4].text = '135~180°'
        # 添加数据
        for i in range(len(side_temp_list)):
            row_cells = table.add_row().cells  # 为每个数据行添加一行
            row_cells[0].text = side_temp_list[i][0]
            row_cells[1].text = side_temp_list[i][1]
            row_cells[2].text = side_temp_list[i][2]
            row_cells[3].text = side_temp_list[i][3]
            row_cells[4].text = side_temp_list[i][4]

        doc.add_heading('2.最大前屈角度信息', level=3)
        # 角度信息	最大角度	最大角速度
        # 外侧肩(就是前端网页选择的那个左侧还是右侧,只需要这个就行)		
        # 创建表格，2行3列（只显示选择的肩部）
        table = doc.add_table(rows=2, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '角度信息'
        hdr_cells[1].text = '最大角度'
        hdr_cells[2].text = '最大角速度'
        # 选择的肩部数据
        selected_cells = table.rows[1].cells
        selected_cells[0].text = shoulder_name
        selected_cells[1].text = str(selected_shoulder_data['max_angle'])
        selected_cells[2].text = str(selected_shoulder_data['max_angle_speed'])

        doc.add_heading('3.最大前屈角角度视图', level=3)
        # 插入最大前屈角角度视图
        if report_data['image_path']['max_flexion_angle'] and os.path.exists(report_data['image_path']['max_flexion_angle']):
            try:
                doc.add_picture(report_data['image_path']['max_flexion_angle'], width=Inches(3))
            except:
                doc.add_paragraph('最大前屈角角度视图（图片加载失败）')
        else:
            doc.add_paragraph('最大前屈角角度视图（图片未生成）')
        doc.add_heading('4.角度和角度速度曲线图', level=3)
        # 插入图片,先不用操作
                # 添加图片
        try:
            doc.add_picture(report_data['image_path']['side_shoulder_angle_speed'], width=Inches(6))
        except:
            doc.add_paragraph('侧面-肩部角度曲线图（图片加载失败）')

        # 直接插入分页符
        doc.add_page_break()


        doc.add_heading('四、后视腕部运动信息', level=2)
        doc.add_heading('1.左右腕部最大高度比例信息', level=3)

        # 创建1行2列的表格，分别为左腕高度比和右腕高度比
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '左腕高度比'
        hdr_cells[1].text = '右腕高度比'
        # 添加数据行
        row_cells = table.add_row().cells
        row_cells[0].text = str(report_data['wrist_joint_data']['left_wrist_height_ratio'])
        row_cells[1].text = str(report_data['wrist_joint_data']['right_wrist_height_ratio'])

        doc.add_heading('2.左右腕部最大高度比图', level=3)
        # 插入左右腕部最大高度比图
        if report_data['image_path']['max_wrist_heights'] and os.path.exists(report_data['image_path']['max_wrist_heights']):
            try:
                doc.add_picture(report_data['image_path']['max_wrist_heights'], width=Inches(6))
            except:
                doc.add_paragraph('左右腕部最大高度比图（图片加载失败）')
        else:
            doc.add_paragraph('左右腕部最大高度比图（图片未生成）')

        doc.add_heading('3.左右腕部高度比例变化曲线图', level=3)
        # 插入图片,先不用操作
                # 添加图片
        try:
            doc.add_picture(report_data['image_path']['back_wrist_height'], width=Inches(6))
        except:
            doc.add_paragraph('背面-手腕高度曲线图（图片加载失败）')

        # 直接插入分页符
        # doc.add_page_break()

        # 结果报告    
        doc.add_heading('五、结果报告', level=2)
        # 三行两列
        # 功能评分：
        # 功能报告：
        # 检查日期：
        # 结果报告    
        # 创建一个表格，每个条目一行
        table = doc.add_table(rows=3, cols=2)  # 创建三行两列的表格
        # 功能评分
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '功能评分：'
        hdr_cells[1].text = str(report_data['result_report']['function_score'])
        # 功能报告
        hdr_cells = table.rows[1].cells
        hdr_cells[0].text = '功能报告：'
        hdr_cells[1].text = report_data['result_report']['function_report']
        # 检查日期
        hdr_cells = table.rows[2].cells
        hdr_cells[0].text = '检查日期：'
        hdr_cells[1].text = report_data['result_report']['check_date']
        
        # 保存文档
        print("报告已生成")
        # 文档名称，报告-姓名-当前信息
        time_str = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        file_name = "report-" + report_data['base_info']['name'] + "-" + time_str + ".docx"
        file_path = os.path.join(output_dir, file_name)
        doc.save(file_path)
        return file_path, file_name 